"""
RELATIONAL LYSIS — MASTER HEADER UPDATE & VALIDATION REPORT GENERATOR
-----------------------------------------------------------------------
Author: Dr. David Swanagon

Functionality:
    1) Updates Python headers:
            - Script Name: file name
            - Category: parent folder name
    2) Executes each script
    3) Captures terminal output
    4) Generates sequential Word report
    5) Overwrites existing output file
    6) Uses active Python interpreter (prevents NumPy errors)
"""

import os
import sys
import subprocess
from docx import Document
from docx.shared import Pt

# ------------------------------------------------------------------
# CONFIGURATION
# ------------------------------------------------------------------

BASE_PATH = "/Users/dswanagon/Relational Lysis Github"

OUTPUT_FILE = os.path.join(
    BASE_PATH,
    "Relational Lysis_Data Validation Experiments.docx"
)

FOLDERS = [
    "1-Core Validation",
    "2-Structural Defense",
    "3-Quantum Reconsitution",
    "4-Topology and Global Validation",
    "5-Internal Mechanism Validation",
    "6-Universal Validation"
]

# ------------------------------------------------------------------
# HEADER UPDATE
# ------------------------------------------------------------------

def update_header(file_path):
    file_name = os.path.basename(file_path)
    folder_name = os.path.basename(os.path.dirname(file_path))

    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    updated_lines = []
    for line in lines:
        stripped = line.strip()

        if stripped.startswith("Script Name:"):
            updated_lines.append(f"Script Name:     {file_name}\n")

        elif stripped.startswith("Category:"):
            updated_lines.append(f"Category:        {folder_name}\n")

        else:
            updated_lines.append(line)

    with open(file_path, "w", encoding="utf-8") as f:
        f.writelines(updated_lines)

# ------------------------------------------------------------------
# SCRIPT EXECUTION (CRITICAL FIX: sys.executable)
# ------------------------------------------------------------------

def run_script(file_path):
    try:
        result = subprocess.run(
            [sys.executable, file_path],   # uses current environment
            capture_output=True,
            text=True,
            timeout=300
        )
        return result.stdout + "\n" + result.stderr

    except Exception as e:
        return f"Execution Error:\n{str(e)}"

# ------------------------------------------------------------------
# WORD FORMATTING
# ------------------------------------------------------------------

def add_code_block(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)

# ------------------------------------------------------------------
# MAIN REPORT GENERATION
# ------------------------------------------------------------------

def generate_report():

    print("Using interpreter:", sys.executable)
    print("Starting validation sweep...\n")

    # Overwrite existing file
    if os.path.exists(OUTPUT_FILE):
        os.remove(OUTPUT_FILE)

    doc = Document()
    doc.add_heading("Relational Lysis_Data Validation Experiments", level=1)

    total_scripts = 0

    for folder in FOLDERS:

        folder_path = os.path.join(BASE_PATH, folder)

        if not os.path.exists(folder_path):
            print(f"Skipping missing folder: {folder_path}")
            continue

        for root, dirs, files in os.walk(folder_path):

            for file in sorted(files):

                if file.endswith(".py"):

                    total_scripts += 1
                    file_path = os.path.join(root, file)

                    print(f"Processing ({total_scripts}): {file_path}")

                    # 1) Update header
                    update_header(file_path)

                    # 2) Read updated script
                    with open(file_path, "r", encoding="utf-8") as f:
                        script_text = f.read()

                    # 3) Execute script
                    terminal_output = run_script(file_path)

                    # 4) Add to Word document
                    doc.add_page_break()
                    doc.add_heading(file, level=2)

                    doc.add_heading("Updated Python Script", level=3)
                    add_code_block(doc, script_text)

                    doc.add_heading("Terminal Results", level=3)
                    add_code_block(doc, terminal_output)

    doc.save(OUTPUT_FILE)

    print("\n-----------------------------------")
    print("VALIDATION COMPLETE")
    print("Total scripts processed:", total_scripts)
    print("Report saved to:")
    print(OUTPUT_FILE)
    print("-----------------------------------")

# ------------------------------------------------------------------
# ENTRY POINT
# ------------------------------------------------------------------

if __name__ == "__main__":
    generate_report()
